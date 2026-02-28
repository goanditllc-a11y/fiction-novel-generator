"""File handler supporting .txt, .docx, and .md formats."""
import os

SUPPORTED = {'.txt', '.docx', '.md'}


class UnsupportedFormatError(Exception):
    pass


class FileHandler:
    """Load and save novel text in multiple formats."""

    # ------------------------------------------------------------------
    def load_file(self, path: str) -> str:
        """Return plain-text content of *path*."""
        ext = os.path.splitext(path)[1].lower()
        if ext not in SUPPORTED:
            raise UnsupportedFormatError(f"Unsupported file extension: {ext}")

        if ext == '.docx':
            return self._load_docx(path)
        # .txt and .md: plain text
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()

    # ------------------------------------------------------------------
    def save_file(self, path: str, content: str) -> None:
        """Save *content* to *path* in the appropriate format."""
        ext = os.path.splitext(path)[1].lower()
        if ext not in SUPPORTED:
            raise UnsupportedFormatError(f"Unsupported file extension: {ext}")

        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)

        if ext == '.docx':
            self._save_docx(path, content)
        else:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(content)

    # ------------------------------------------------------------------
    def _load_docx(self, path: str) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required to open .docx files.")
        doc = Document(path)
        paragraphs = [para.text for para in doc.paragraphs]
        return '\n'.join(paragraphs)

    def _save_docx(self, path: str, content: str) -> None:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required to save .docx files.")
        doc = Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(path)
