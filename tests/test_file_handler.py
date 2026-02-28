"""Tests for FileHandler."""
import sys, os, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

import pytest
from src.file_io.file_handler import FileHandler, UnsupportedFormatError


@pytest.fixture
def handler():
    return FileHandler()


SAMPLE_TEXT = "Once upon a time, there was a brave knight.\nHe rode into the dark forest."


def test_load_txt_file(handler, tmp_path):
    p = tmp_path / "novel.txt"
    p.write_text(SAMPLE_TEXT, encoding='utf-8')
    loaded = handler.load_file(str(p))
    assert loaded == SAMPLE_TEXT


def test_save_txt_file(handler, tmp_path):
    p = tmp_path / "output.txt"
    handler.save_file(str(p), SAMPLE_TEXT)
    assert p.read_text(encoding='utf-8') == SAMPLE_TEXT


def test_save_and_reload_txt(handler, tmp_path):
    p = tmp_path / "roundtrip.txt"
    handler.save_file(str(p), SAMPLE_TEXT)
    loaded = handler.load_file(str(p))
    assert loaded == SAMPLE_TEXT


def test_load_md_file(handler, tmp_path):
    p = tmp_path / "novel.md"
    content = "# Chapter 1\n\nOnce upon a time."
    p.write_text(content, encoding='utf-8')
    loaded = handler.load_file(str(p))
    assert loaded == content


def test_save_md_file(handler, tmp_path):
    p = tmp_path / "novel.md"
    content = "# My Novel\n\nParagraph one."
    handler.save_file(str(p), content)
    assert p.read_text(encoding='utf-8') == content


def test_unsupported_extension_raises(handler, tmp_path):
    p = tmp_path / "novel.odt"
    with pytest.raises(UnsupportedFormatError):
        handler.load_file(str(p))


def test_unsupported_save_extension_raises(handler, tmp_path):
    p = tmp_path / "novel.pdf"
    with pytest.raises(UnsupportedFormatError):
        handler.save_file(str(p), "text")


def test_load_docx_file(handler, tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    p = tmp_path / "novel.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(p))
    loaded = handler.load_file(str(p))
    assert "First paragraph." in loaded
    assert "Second paragraph." in loaded
